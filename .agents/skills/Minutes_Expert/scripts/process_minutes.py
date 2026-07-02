# -*- coding: utf-8 -*-
"""
회의록 자동 작성을 처리하는 핵심 파이썬 스크립트입니다.
"""

import os # OS 관련 기능을 사용하기 위한 모듈
import docx # docx 문서를 읽고 쓰기 위한 모듈
from docx.oxml import OxmlElement # XML 엘리먼트 조작을 위한 모듈
from docx.text.paragraph import Paragraph # 문단 객체 조작을 위한 모듈

def parse_sample_minutes(sample_path):
    # 샘플 회의록 문서를 읽어 필요한 데이터를 추출하는 함수
    doc = docx.Document(sample_path) # 샘플 회의록 문서 객체 생성
    
    data = { # 파싱된 회의록 데이터를 구조화하여 저장할 딕셔너리
        '일시': '', # 회의 일시를 저장할 문자열
        '장소': '', # 회의 장소를 저장할 문자열
        '참석자': '', # 회의 참석자들을 저장할 문자열
        'agendas': [], # 핵심 안건 목록을 저장할 리스트
        'discussions': [], # 주요 논의 내용을 저장할 리스트
        'decisions': {}, # 결정 사항들을 저장할 딕셔너리
        'next_meeting': '', # 다음 회의 일정을 저장할 문자열
        'writer': '', # 회의록 작성자명을 저장할 문자열
        'table_data': [] # 향후 과제 테이블 데이터를 저장할 리스트
    }
    
    current_section = None # 현재 분석 중인 회의록 섹션을 저장할 상태 변수
    
    for para in doc.paragraphs: # para: 샘플 회의록 문서의 각 문단 객체
        text = para.text.strip() # 문단의 텍스트 공백 제거 문자열
        if not text: # 빈 문단인 경우 건너뛰기
            continue
            
        if text.startswith('---'): # 구분선 문단인 경우 건너뛰기
            continue
            
        # 섹션 변경 감지
        if "📅 회의 정보" in text: # 회의 정보 섹션 시작 감지
            current_section = "info" # 상태를 회의 정보로 변경
            continue
        elif "📝 핵심 안건" in text: # 핵심 안건 섹션 시작 감지
            current_section = "agenda" # 상태를 핵심 안건으로 변경
            continue
        elif "💬 주요 논의 내용" in text: # 주요 논의 내용 섹션 시작 감지
            current_section = "discussion" # 상태를 주요 논의 내용으로 변경
            continue
        elif "✅ 결정 사항" in text: # 결정 사항 섹션 시작 감지
            current_section = "decision" # 상태를 결정 사항으로 변경
            continue
        elif "🚀 향후 과제" in text: # 향후 과제 섹션 시작 감지
            current_section = "action" # 상태를 향후 과제로 변경
            continue
        elif "📎 기타 참고 사항" in text: # 기타 참고 사항 섹션 시작 감지
            current_section = "etc" # 상태를 기타 참고 사항으로 변경
            continue
            
        # 현재 섹션에 따라 알맞게 데이터 파싱
        if current_section == "info": # 회의 정보 파싱 상태인 경우
            if text.startswith("일시:"): # 일시 데이터인 경우
                data['일시'] = text.replace("일시:", "").strip() # 일시 값 추출 및 저장
            elif text.startswith("장소:"): # 장소 데이터인 경우
                data['장소'] = text.replace("장소:", "").strip() # 장소 값 추출 및 저장
            elif text.startswith("참석자:"): # 참석자 데이터인 경우
                data['참석자'] = text.replace("참석자:", "").strip() # 참석자 값 추출 및 저장
        elif current_section == "agenda": # 핵심 안건 파싱 상태인 경우
            data['agendas'].append(text) # 안건 텍스트 추가
        elif current_section == "discussion": # 주요 논의 내용 파싱 상태인 경우
            data['discussions'].append(text) # 논의 내용 텍스트 추가
        elif current_section == "decision": # 결정 사항 파싱 상태인 경우
            if text.startswith("[결정"): # 결정 사항 데이터인 경우
                parts = text.split("]", 1) # 결정 번호와 상세 내용을 분리한 리스트
                if len(parts) == 2: # 성공적으로 분리된 경우
                    key = parts[0] + "]" # 결정 키 문자열 (예: [결정 01])
                    val = parts[1].strip() # 결정 내용 상세 문자열
                    data['decisions'][key] = val # 결정 사항 딕셔너리에 추가
        elif current_section == "etc": # 기타 참고 사항 파싱 상태인 경우
            if "다음 회의 일정:" in text: # 다음 회의 일정 텍스트가 포함된 경우
                val = text.split("다음 회의 일정:")[1].strip() # 다음 회의 일정 값 추출
                data['next_meeting'] = val # 딕셔너리에 저장
            elif "회의록 작성자:" in text: # 회의록 작성자 텍스트가 포함된 경우
                val = text.split("회의록 작성자:")[1].strip() # 작성자명 값 추출
                data['writer'] = val # 딕셔너리에 저장

    if len(doc.tables) > 0: # 문서에 테이블이 존재하는 경우
        table = doc.tables[0] # 첫 번째 테이블 객체 선택
        for row in table.rows[1:]: # row: 테이블의 헤더를 제외한 각 데이터 행 객체
            row_cells = [cell.text.strip() for cell in row.cells] # 행 내부 각 셀의 공백 제거 텍스트 리스트
            if any(row_cells): # 데이터가 하나라도 들어있는 행인 경우
                data['table_data'].append(row_cells) # 테이블 데이터 목록에 추가
                
    return data # 최종 파싱된 회의록 데이터 딕셔너리 반환

def insert_paragraph_after(paragraph, text=None, style=None):
    # docx 문서 내 지정한 문단 뒤에 새로운 문단을 삽입하는 헬퍼 함수
    new_p = OxmlElement('w:p') # 새로운 w:p XML 엘리먼트 생성
    paragraph._p.addnext(new_p) # XML 노드상 현재 문단 뒤에 새 문단 엘리먼트 추가
    new_paragraph = Paragraph(new_p, paragraph._parent) # Paragraph 객체로 래핑
    if text: # 텍스트가 입력된 경우
        new_paragraph.text = text # 문단 텍스트 설정
    if style: # 스타일이 지정된 경우
        new_paragraph.style = style # 문단 스타일 설정
    return new_paragraph # 생성된 신규 문단 객체 반환

def fill_minutes_template(template_path, output_path, data):
    # 파싱된 데이터를 템플릿에 매핑하여 최종 회의록 문서를 생성하는 함수
    doc = docx.Document(template_path) # 템플릿 docx 파일 로드 및 문서 객체 생성
    
    agenda_heading_para = None # 핵심 안건 헤더 문단의 객체를 임시 저장할 변수
    discussion_heading_para = None # 주요 논의 내용 헤더 문단의 객체를 임시 저장할 변수
    
    # 1. 템플릿 문단들을 순회하며 기본 텍스트 치환 및 타겟 헤더 위치 검색
    for para in doc.paragraphs: # para: 템플릿 문서의 각 문단 객체
        text = para.text.strip() # 문단 텍스트의 양 끝 공백 제거 문자열
        
        # 회의 정보 채우기
        if text.startswith("일시:"): # 일시 입력 필드인 경우
            para.text = f"일시: {data['일시']}" # 일시 데이터 입력
        elif text.startswith("장소:"): # 장소 입력 필드인 경우
            para.text = f"장소: {data['장소']}" # 장소 데이터 입력
        elif text.startswith("참석자:"): # 참석자 입력 필드인 경우
            para.text = f"참석자: {data['참석자']}" # 참석자 데이터 입력
            
        # 핵심 안건 헤더 위치 기록
        elif "📝 핵심 안건" in text: # 핵심 안건 섹션 헤더인 경우
            agenda_heading_para = para # 헤더 문단 객체 저장
            
        # 주요 논의 내용 헤더 위치 기록
        elif "💬 주요 논의 내용" in text: # 주요 논의 내용 섹션 헤더인 경우
            discussion_heading_para = para # 헤더 문단 객체 저장
            
        # 결정 사항 업데이트
        elif text.startswith("[결정"): # 결정 사항 문단인 경우
            for key, val in data['decisions'].items(): # key: 결정 번호(예: [결정 01]), val: 결정 내용
                if text.startswith(key): # 결정 번호 키가 매칭되는 경우
                    para.text = f"{key} {val}" # 해당 결정 내용으로 업데이트
                    break
                    
        # 기타 참고 사항 채우기
        elif text.startswith("• 다음 회의 일정:"): # 다음 일정 입력 필드인 경우
            para.text = f"• 다음 회의 일정: {data['next_meeting']}" # 다음 일정 데이터 입력
        elif text.startswith("• 회의록 작성자:"): # 작성자 입력 필드인 경우
            para.text = f"• 회의록 작성자: {data['writer']}" # 작성자 데이터 입력
            
    # 2. 동적으로 늘어날 수 있는 문단들 삽입 처리
    # 핵심 안건 항목 동적 삽입
    if agenda_heading_para and data['agendas']: # 안건 헤더 문단과 안건 데이터가 존재하는 경우
        current_p = agenda_heading_para # 삽입 기준 문단 변수 초기화
        for agenda in data['agendas']: # agenda: 삽입할 각 안건 텍스트
            current_p = insert_paragraph_after(current_p, text=agenda, style='Normal') # 안건 문단을 순차적으로 삽입
            
    # 주요 논의 내용 항목 동적 삽입
    if discussion_heading_para and data['discussions']: # 논의 내용 헤더 문단과 논의 데이터가 존재하는 경우
        current_p = discussion_heading_para # 삽입 기준 문단 변수 초기화
        for discussion in data['discussions']: # discussion: 삽입할 각 논의 내용 텍스트
            current_p = insert_paragraph_after(current_p, text=discussion, style='Normal') # 논의 내용 문단을 순차적으로 삽입
            
    # 3. 테이블 데이터 업데이트
    if len(doc.tables) > 0 and len(data['table_data']) > 0: # 문서에 테이블 및 업데이트할 데이터가 있는 경우
        table = doc.tables[0] # 첫 번째 테이블 객체 선택
        total_rows = len(table.rows) # 현재 테이블의 총 행 수
        needed_rows = 1 + len(data['table_data']) # 필요한 행 수 (헤더 1행 + 데이터 행 수)
        
        # 템플릿에 정의된 여분의 빈 행 삭제
        if total_rows > needed_rows: # 행 수가 과다할 경우
            for _ in range(total_rows - needed_rows): # 초과하는 행 수만큼 반복
                tbl = table._tbl # XML 테이블 객체
                tbl.remove(table.rows[-1]._tr) # 마지막 행 삭제
        # 행이 부족할 경우 추가
        elif total_rows < needed_rows: # 행 수가 부족할 경우
            for _ in range(needed_rows - total_rows): # 부족한 행 수만큼 반복
                table.add_row() # 신규 행 추가
                
        # 데이터 매핑 입력
        for row_idx, row_data in enumerate(data['table_data']): # row_idx: 데이터 인덱스, row_data: 각 행의 데이터 리스트
            row = table.rows[row_idx + 1] # 헤더 다음 행부터 순차적 선택
            for col_idx, cell_value in enumerate(row_data): # col_idx: 열 인덱스, cell_value: 입력값 문자열
                if col_idx < len(row.cells): # 열 개수 범위를 초과하지 않는 경우
                    row.cells[col_idx].text = cell_value # 셀 데이터 입력
                    
    # 최종 결과 파일 저장
    doc.save(output_path) # 완성된 회의록 문서 저장

# 실행 영역
script_dir = os.path.dirname(os.path.abspath(__file__)) # 현재 실행 중인 스크립트 파일의 절대 경로
template_path = os.path.join(script_dir, "../resources/minutes_template.docx") # 템플릿 파일의 절대 경로 변환
sample_path = os.path.join(script_dir, "../../../../실습1-1/minutes_sample.docx") # 샘플 회의록 파일의 절대 경로 변환
output_path = os.path.join(script_dir, "../../../../실습1-1/output_minutes.docx") # 결과 파일 저장 절대 경로 변환

if __name__ == "__main__":
    print("회의록 작성을 시작합니다...") # 시작 메시지 출력
    print(f"템플릿 경로: {template_path}") # 사용 템플릿 경로 콘솔 출력
    print(f"샘플 회의록 경로: {sample_path}") # 원본 샘플 파일 경로 콘솔 출력
    
    parsed_data = parse_sample_minutes(sample_path) # 샘플 파일 파싱 실행 및 데이터 저장
    fill_minutes_template(template_path, output_path, parsed_data) # 데이터 채우기 및 파일 저장 실행
    
    print(f"회의록 생성이 완료되었습니다. 저장 경로: {output_path}") # 완료 메시지 출력
