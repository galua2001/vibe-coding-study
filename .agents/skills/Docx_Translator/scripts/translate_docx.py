# -*- coding: utf-8 -*-
"""
docx 워드 문서 내부의 텍스트를 감지하여 한글로 자동 번역하는 핵심 파이썬 스크립트입니다.
"""

import os # OS 관련 기능을 제어하기 위한 모듈
import docx # docx 워드 문서를 읽고 쓰기 위한 모듈
from deep_translator import GoogleTranslator # 초고속 구글 번역 모듈 로드

def create_sample_english_docx(file_path):
    # 테스트용 영문 docx 파일이 없을 경우 예시 파일을 자동으로 생성하는 함수
    doc = docx.Document() # 빈 문서 객체 생성
    doc.add_heading('Antigravity Project Proposal', level=1) # 1단계 제목 추가
    
    p1 = doc.add_paragraph('This is a project proposal for our new AI assistant platform.') # 문단 1 추가
    p1.add_run(' We aim to maximize productivity.').bold = True # 굵은 글씨 추가
    
    doc.add_paragraph('Key Features:') # 문단 2 추가
    doc.add_paragraph('- Auto coding with multi-agents') # 목록 1 추가
    doc.add_paragraph('- Secure version control and environment setups') # 목록 2 추가
    
    # 2행 3열짜리 간단한 영어 표 생성
    table = doc.add_table(rows=2, cols=3) # 표 객체 생성
    hdr_cells = table.rows[0].cells # 헤더 셀 목록
    hdr_cells[0].text = 'Task Name' # 1열 헤더
    hdr_cells[1].text = 'Owner' # 2열 헤더
    hdr_cells[2].text = 'Deadline' # 3열 헤더
    
    row_cells = table.rows[1].cells # 데이터 행 1의 셀 목록
    row_cells[0].text = 'Setup python-docx skill' # 1열 데이터
    row_cells[1].text = 'Antigravity AI' # 2열 데이터
    row_cells[2].text = 'Today' # 3열 데이터
    
    doc.save(file_path) # 지정된 경로에 문서 저장
    print(f"[시스템] 테스트용 영어 샘플 파일({file_path})을 생성했습니다.")

def translate_docx_file(input_path, output_path):
    # 영문 docx 파일의 서식을 최대한 유지하면서 내용을 한국어로 번역하는 함수
    doc = docx.Document(input_path) # 원본 docx 문서 객체 로드
    translator = GoogleTranslator(source='auto', target='ko') # 번역기 객체 생성 (언어 자동감지 ➡️ 한국어)
    
    # 1. 문서 일반 문단(Paragraphs) 번역 진행
    for para in doc.paragraphs: # para: 문서 내의 각 문단 객체
        if not para.text.strip(): # 빈 문단인 경우 건너뛰기
            continue
            
        # 문단 내부의 스타일(볼드, 이탤릭, 폰트 등)을 보존하기 위해 run 단위로 쪼개어 번역
        for run in para.runs: # run: 문단 내 동일 서식을 가진 최소 텍스트 묶음 객체
            original_text = run.text.strip() # 원본 텍스트 공백 제거
            if original_text and len(original_text) > 1: # 번역할 텍스트가 존재하는 경우
                translated_text = translator.translate(original_text) # 구글 번역 수행
                run.text = run.text.replace(original_text, translated_text) # 원본 단어를 번역된 단어로 교체
                
    # 2. 문서 내부 표(Tables) 데이터 번역 진행
    for table in doc.tables: # table: 문서 내의 각 표 객체
        for row in table.rows: # row: 표의 각 행 객체
            for cell in row.cells: # cell: 행의 각 열 셀 객체
                for para in cell.paragraphs: # 셀 내부 문단 순회
                    for run in para.runs: # 셀 내 문단 서식 텍스트 순회
                        original_text = run.text.strip() # 원본 텍스트
                        if original_text and len(original_text) > 1: # 번역이 필요한 경우
                            translated = translator.translate(original_text) # 번역 수행
                            run.text = run.text.replace(original_text, translated) # 번역본으로 치환
                            
    doc.save(output_path) # 번역이 완료된 문서 저장
    print(f"[성공] 번역된 한글 문서를 저장했습니다: {output_path}")

# 실행 영역
script_dir = os.path.dirname(os.path.abspath(__file__)) # 현재 스크립트의 절대 경로
input_docx = os.path.join(script_dir, "../../../실습1-1/input.docx") # 입력 영어 문서 절대 경로
output_docx = os.path.join(script_dir, "../../../실습1-1/translated_output.docx") # 결과물 한글 문서 절대 경로

if __name__ == "__main__":
    print("문서 자동 번역을 시작합니다...") # 시작 안내 출력
    
    # 만약 입력 파일이 없을 경우, 예시로 쓸 영어 문서를 자동 생성
    if not os.path.exists(input_docx): # 파일이 존재하지 않는 경우
        create_sample_english_docx(input_docx) # 영어 샘플 문서 생성 함수 실행
        
    translate_docx_file(input_docx, output_docx) # 번역 및 파일 저장 실행
    print("모든 번역 작업이 완료되었습니다!") # 완료 메시지 출력
